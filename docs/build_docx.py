"""
Eggplant 홍보 보고서를 Word 파일(.docx)로 변환.

사용:
    python3 docs/build_docx.py

생성:
    docs/Eggplant_홍보보고서.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─────────────────────────────────────────────────────────────────────
# 색상 정의
PURPLE = RGBColor(0x93, 0x33, 0xEA)        # eggplant primary
DARK_PURPLE = RGBColor(0x7E, 0x22, 0xCE)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
TEXT_DARK = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_background(cell, color_hex):
    """셀 배경색 설정."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    """스타일 적용된 헤딩."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if color is not None:
            run.font.color.rgb = color
    return h


def add_para(doc, text, size=11, bold=False, color=None, italic=False, center=False):
    """본문 단락 추가."""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, indent=0):
    """글머리 기호."""
    p = doc.add_paragraph(style='List Bullet')
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.runs[0] if p.runs else p.add_run()
    p.text = ''
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(11)


def add_quote_box(doc, text, color=PURPLE):
    """인용/강조 박스 (1셀 표 형태)."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_background(cell, '{:02X}{:02X}{:02X}'.format(0xFA, 0xF5, 0xFF))  # eggplant-50
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_simple_table(doc, headers, rows, col_widths=None):
    """심플 표."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = WHITE
        set_cell_background(hdr[i], '9333EA')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Body
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table


def add_pagebreak(doc):
    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# 보고서 빌드
# ═════════════════════════════════════════════════════════════════════
doc = Document()

# 기본 폰트
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = Pt(11)

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


# ═════ 표지 ═════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(120)
run = title_p.add_run('🍆')
run.font.size = Pt(72)

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p2.add_run('Eggplant')
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = PURPLE

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('홍보 보고서')
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = TEXT_DARK

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('익명으로 안전한 동네 중고거래,\n그리고 사용자가 보상받는 마켓.')
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(14)
run.italic = True
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()
ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_p.paragraph_format.space_before = Pt(60)
run = ver_p.add_run('마케팅 / PR / 광고 가이드 v1.0')
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026.04.28')
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

add_pagebreak(doc)


# ═════ 목차 ═════════════════════════════════════════════════════════
add_heading(doc, '📑 목차', level=1, color=PURPLE)
toc = [
    '1. 한 문장 요약 (Elevator Pitch)',
    '2. 왜 지금 Eggplant인가 — 시장 진단',
    '3. 핵심 차별점 5가지',
    '4. QTA 토큰 이코노미',
    '5. 주요 기능 한눈에 보기',
    '6. 기술 스택',
    '7. 타깃 사용자 페르소나',
    '8. 경쟁사 비교',
    '9. 홍보 메시지 카피',
    '10. SNS / 광고 콘텐츠 가이드',
    '11. FAQ — 자주 묻는 질문',
]
for item in toc:
    add_para(doc, item, size=12, bold=True, color=TEXT_DARK)

add_pagebreak(doc)


# ═════ 1. 한 문장 요약 ═════
add_heading(doc, '1. 한 문장 요약 (Elevator Pitch)', level=1, color=PURPLE)
add_quote_box(doc,
    '"전화번호 없이 닉네임으로 가입, QR 한 번이면 거래 시작.\n그리고 사용할수록 QTA 토큰이 쌓입니다."'
)

doc.add_paragraph()
add_heading(doc, '30초 피치', level=2, color=DARK_PURPLE)
add_para(doc,
    '국내 중고거래 시장은 1위 앱이 압도적이지만, 개인정보 노출, 사기, 분쟁 같은 '
    '고질적 문제가 있습니다. Eggplant는 완전 익명 가입(닉네임 + 지갑주소), '
    '휘발성 채팅(메시지 서버 미저장), QR 코드 만남으로 이 문제들을 정면 돌파합니다. '
    '게다가 거래·로그인·친구 초대마다 QTA 토큰이 쌓이고, 5,000 QTA부터 출금 가능 — '
    '사용자가 마켓의 주인이 됩니다.'
)


# ═════ 2. 시장 진단 ═════
add_pagebreak(doc)
add_heading(doc, '2. 왜 지금 Eggplant인가 — 시장 진단', level=1, color=PURPLE)

add_heading(doc, '중고거래 시장의 3대 페인포인트', level=2, color=DARK_PURPLE)
add_simple_table(doc,
    headers=['문제', '실제 사례', 'Eggplant 해법'],
    rows=[
        ['개인정보 노출', '거래 중 전화번호 유출 → 스팸 / 스토킹 우려', '닉네임만 사용, 전화번호 입력란 자체가 없음'],
        ['채팅 기록 영구 저장', '분쟁 시 사적 대화 내역이 남고, 해킹 시 유출', '메시지는 서버 무저장 + 화면 벗어나면 증발'],
        ['신뢰할 수 없는 거래자', '사기·노쇼·잠수 빈번', '매너온도 + QR 만남 + 거래 보너스로 신뢰 형성'],
    ]
)

doc.add_paragraph()
add_heading(doc, '글로벌 트렌드 부합', level=2, color=DARK_PURPLE)
add_bullet(doc, 'Privacy-First 추세 (Apple ATT, GDPR 등)')
add_bullet(doc, 'Web3 / Token Economy 확산 (사용자가 보상받는 플랫폼)')
add_bullet(doc, 'Hyper-local Commerce (동네 기반 직거래) 성장')


# ═════ 3. 핵심 차별점 5가지 ═════
add_pagebreak(doc)
add_heading(doc, '3. 핵심 차별점 5가지', level=1, color=PURPLE)

diffs = [
    ('🔐 ① 완전 익명 — "전화번호? 그게 뭔가요?"', [
        '가입에 필요한 정보: 닉네임 + 지갑주소(0x로 시작) + 비밀번호 단 3가지',
        '전화번호, 이메일, 실명, 생년월일 일체 수집 안 함',
        '비밀번호 분실 시 지갑주소로 복구',
    ]),
    ('💨 ② 휘발성 채팅 (QRChat)', [
        '메시지가 서버 DB에 절대 저장되지 않음 (Durable Object 메모리에만 일시 보관)',
        '채팅방을 나가면 상대방 화면에서도 메시지 사라짐',
        '스크린샷 차단 (Android FLAG_SECURE)',
        '카톡·문자처럼 흔적이 남지 않아 분쟁 후 마음의 짐 X',
    ]),
    ('📱 ③ QR 코드로 만남', [
        '직접 만나서 거래할 때 내 QR 코드만 보여주면 됨',
        '상대가 스캔하면 자동으로 채팅방 / 거래 인증 연결',
        '전화번호 교환 없이 한 번의 거래로 완전 분리',
    ]),
    ('🪙 ④ QTA 토큰 보상 — 쓸수록 돈이 모입니다', [
        '가입 즉시 +500 QTA',
        '매일 로그인 +10 QTA × 최대 3회',
        '거래 완료 시 +10 QTA (구매자/판매자 양쪽)',
        '친구 1명 초대마다 +200 QTA (무제한!)',
        '5,000 QTA부터 출금 가능 (5,000 단위)',
    ]),
    ('🏘️ ⑤ 동네 기반 + 매너온도', [
        '동네 인증 (위치 기반)으로 이웃과만 거래',
        '매너온도(36.5°C 시작) — 후기마다 ±0.5°C 자동 반영',
        '익명이지만 신뢰는 누적',
    ]),
]
for title, items in diffs:
    add_heading(doc, title, level=2, color=DARK_PURPLE)
    for item in items:
        add_bullet(doc, item)


# ═════ 4. QTA 토큰 이코노미 ═════
add_pagebreak(doc)
add_heading(doc, '4. QTA 토큰 이코노미', level=1, color=PURPLE)

add_heading(doc, '💎 적립 구조', level=2, color=DARK_PURPLE)
add_simple_table(doc,
    headers=['항목', '보상', '조건'],
    rows=[
        ['🎁 신규 가입', '+500 QTA', '1회, 즉시'],
        ['📅 매일 로그인', '+10 QTA × 최대 3회/일', 'UTC 자정 리셋'],
        ['🤝 거래 완료', '+10 QTA', '구매·판매 양쪽 모두'],
        ['👥 친구 초대', '+200 QTA', '1명당, 무제한'],
        ['💰 상품 결제 (QTA)', '즉시 자동 송금', '구매자 → 판매자, 멱등 처리'],
    ]
)

doc.add_paragraph()
add_heading(doc, '💸 출금', level=2, color=DARK_PURPLE)
add_bullet(doc, '최소 5,000 QTA부터 신청 가능')
add_bullet(doc, '5,000 단위로만 출금')
add_bullet(doc, '운영자 검토 후 처리 (수동 승인 → 자동 송금)')

doc.add_paragraph()
add_heading(doc, '🎯 토큰 이코노미가 주는 사용자 경험', level=2, color=DARK_PURPLE)
add_bullet(doc, '"어차피 중고거래 할 거, 돈이 쌓이는 앱에서 하자"')
add_bullet(doc, '친구 초대로 소셜 그로스 자체적 동력')
add_bullet(doc, '매일 들어오면 보너스 → DAU 자연 증가')
add_bullet(doc, '거래 완료 인센티브로 노쇼·잠수 감소')


# ═════ 5. 주요 기능 ═════
add_pagebreak(doc)
add_heading(doc, '5. 주요 기능 한눈에 보기', level=1, color=PURPLE)

add_heading(doc, '🛒 거래 핵심 기능', level=2, color=DARK_PURPLE)
add_simple_table(doc,
    headers=['기능', '설명'],
    rows=[
        ['상품 등록', '사진 최대 10장 + 영상 1개 + 설명 + 가격(KRW or QTA)'],
        ['카테고리 / 지역 필터', '동네 기반으로 빠르게 매물 탐색'],
        ['가격 제안', '구매 희망자가 원하는 가격으로 협상'],
        ['상품 끌어올리기', '노출 우선순위 향상 (24시간 쿨타임)'],
        ['QTA 결제', '상품 가격을 QTA로 책정 시 자동 토큰 송금'],
        ['거래 후 후기', '매너 좋아요 / 그저 그래요 / 별로예요 (3단계)'],
    ]
)

doc.add_paragraph()
add_heading(doc, '💬 커뮤니케이션', level=2, color=DARK_PURPLE)
add_simple_table(doc,
    headers=['기능', '설명'],
    rows=[
        ['휘발성 채팅', '1:1 실시간 메시지, 서버 무저장'],
        ['WebRTC P2P 음성통화', '채팅방 내 음성 통화 (서버 경유 X)'],
        ['키워드 알림', '원하는 상품 키워드 등록 → 매물 등록 시 푸시'],
        ['숨김 / 신고 / 차단', '불쾌한 상품·사용자 즉시 차단'],
    ]
)

doc.add_paragraph()
add_heading(doc, '👤 프로필 / 신원', level=2, color=DARK_PURPLE)
add_simple_table(doc,
    headers=['기능', '설명'],
    rows=[
        ['매너온도', '36.5°C 시작, 후기마다 변동'],
        ['거래 후기 보기', '다른 사용자의 거래 평판 확인'],
        ['동네 인증', '위치 기반 동네 설정'],
        ['닉네임으로 비번 복구', '지갑주소 입력 → 새 비밀번호 설정'],
    ]
)

doc.add_paragraph()
add_heading(doc, '🛡️ 안전 기능', level=2, color=DARK_PURPLE)
add_simple_table(doc,
    headers=['기능', '설명'],
    rows=[
        ['스크린샷 차단', '채팅 화면 캡처 자체가 안 됨 (Android)'],
        ['본인 기기 잠금', '새 기기에서 로그인 시 기존 기기 토큰 무효화'],
        ['알림 본문 마스킹', '잠금화면에 메시지 내용 안 보이도록 토글'],
        ['계정 완전 삭제', '데이터 영구 삭제 (탈퇴 시 즉시)'],
    ]
)


# ═════ 6. 기술 스택 ═════
add_pagebreak(doc)
add_heading(doc, '6. 기술 스택', level=1, color=PURPLE)

add_heading(doc, '📱 모바일 앱 (Flutter)', level=2, color=DARK_PURPLE)
add_bullet(doc, 'Flutter 3.22+ / Dart — Android·iOS 동시 지원')
add_bullet(doc, 'WebRTC — 서버 거치지 않는 P2P 음성 통화')
add_bullet(doc, 'Mobile Scanner — 빠른 QR 인식')
add_bullet(doc, '반응형 UI — 폴드폰·태블릿까지 자동 대응 (max-width 600dp)')

add_heading(doc, '☁️ 백엔드 (Cloudflare Workers)', level=2, color=DARK_PURPLE)
add_bullet(doc, 'Cloudflare Workers — 글로벌 엣지에서 5초 미만 콜드스타트')
add_bullet(doc, 'D1 (SQLite) — 사용자·상품·거래 데이터')
add_bullet(doc, 'R2 — 상품 이미지·영상 저장 (S3 호환, 99.99% 가용성)')
add_bullet(doc, 'Durable Objects — 채팅·시그널링 (메시지 무저장)')
add_bullet(doc, 'JWT (HS256) — 가벼운 토큰 인증')

add_heading(doc, '🔒 보안 / 프라이버시', level=2, color=DARK_PURPLE)
add_bullet(doc, 'PBKDF2-SHA256 (100,000 iter) — 비밀번호 해싱 (OWASP 권장)')
add_bullet(doc, 'Pinpoint Salt — 사용자별 무작위 솔트')
add_bullet(doc, 'Token Version — 기기 변경 시 이전 토큰 자동 무효화')
add_bullet(doc, 'Rate Limit — 닉네임 검색 30회/분/IP')


# ═════ 7. 페르소나 ═════
add_pagebreak(doc)
add_heading(doc, '7. 타깃 사용자 페르소나', level=1, color=PURPLE)

personas = [
    ('🌸 페르소나 A — "프라이버시를 중시하는 직장인"',
     '30대 여성 / 서울 강남구 거주 / 회사원',
     '"당근에 연락처 들어가는 게 너무 불편해요. 거래 후에 광고 문자가 와요."',
     ['익명으로 거래 → 사후 스팸 0건', '휘발성 채팅 → 거래 후 흔적 0건']),
    ('🎓 페르소나 B — "용돈이 부족한 대학생"',
     '20대 남성 / 대학가 자취 / 월 60만 원 생활비',
     '"어차피 중고거래 자주 해. 그럴 거면 돈 쌓이는 앱이 낫지."',
     ['가입만 해도 +500 QTA', '친구 초대 +200 QTA × 무제한 → 학과 단톡방 1방이면 즉시 출금각']),
    ('👨‍👩‍👧 페르소나 C — "안전이 우선인 부모"',
     '40대 여성 / 학부모 / 동네 거래 위주',
     '"애들 옷, 장난감 자주 거래하는데 모르는 사람 만나는 게 늘 불안해요."',
     ['QR 만남 → 만남 시간만 짧게 노출', '매너온도로 사전 신뢰 확인',
      '동네 인증된 사용자만 보임']),
    ('💻 페르소나 D — "Web3에 관심 있는 얼리어답터"',
     '20-30대 / IT 종사자 / 코인·NFT 관심층',
     '"토큰 이코노미가 결합된 일상 서비스 사례를 찾고 있어요."',
     ['QTA 토큰 자동 적립', '지갑주소 기반 인증 (Web3 친화)', '출금 가능한 실용적 토큰']),
]
for title, profile, quote, points in personas:
    add_heading(doc, title, level=2, color=DARK_PURPLE)
    add_para(doc, profile, size=10, italic=True, color=GRAY)
    add_quote_box(doc, quote, color=DARK_PURPLE)
    add_para(doc, '✅ Eggplant 매력 포인트:', size=11, bold=True, color=GREEN)
    for p in points:
        add_bullet(doc, p)
    doc.add_paragraph()


# ═════ 8. 경쟁사 비교 ═════
add_pagebreak(doc)
add_heading(doc, '8. 경쟁사 비교', level=1, color=PURPLE)
add_simple_table(doc,
    headers=['항목', '🍆 Eggplant', '🥕 당근마켓', '🐼 번개장터', '🛍️ 중고나라'],
    rows=[
        ['익명성', '✅ 닉네임만', '❌ 전화번호', '❌ 전화번호', '❌ 회원가입'],
        ['채팅 기록', '✅ 휘발성', '❌ 영구 저장', '❌ 영구 저장', '🟡 카페 글'],
        ['토큰 보상', '✅ QTA 적립·출금', '❌ 없음', '🟡 일부', '❌ 없음'],
        ['QR 만남', '✅ 기본 기능', '❌ 없음', '❌ 없음', '❌ 없음'],
        ['스크린샷 차단', '✅ 채팅 화면', '❌ 가능', '❌ 가능', '❌ 가능'],
        ['P2P 통화', '✅ WebRTC', '🟡 채팅만', '🟡 채팅만', '❌'],
        ['친구 초대 보상', '✅ +200 QTA × ∞', '❌', '🟡 일부', '❌'],
        ['매너온도', '✅ 36.5°C', '✅ 동일', '❌', '❌'],
        ['글로벌 인프라', '✅ CF 엣지', '🟡 국내', '🟡 국내', '❌ 카페'],
    ]
)
doc.add_paragraph()
add_quote_box(doc,
    '🎯 결론: Eggplant는 "익명 + 휘발성 + 토큰" 3박자를 모두 갖춘 유일한 중고거래 앱입니다.'
)


# ═════ 9. 홍보 카피 ═════
add_pagebreak(doc)
add_heading(doc, '9. 홍보 메시지 카피', level=1, color=PURPLE)

add_heading(doc, '🎯 메인 슬로건 후보', level=2, color=DARK_PURPLE)
slogans = [
    '"전화번호 없이, 흔적 없이, 그리고 보상받으며."',
    '"가지처럼 익명으로, QR로 만나는 동네 마켓."',
    '"중고거래, 이제 사용자가 주인입니다." 🍆 → QTA로 환원',
    '"닉네임 하나로 시작하는 가장 안전한 거래."',
    '"채팅이 사라지는 마켓 — 당신의 프라이버시는 영원합니다."',
]
for i, s in enumerate(slogans, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = PURPLE
    run = p.add_run(s)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, '💬 짧은 카피 (소셜용)', level=2, color=DARK_PURPLE)

short_copies = [
    """🍆 가입 보너스 +500 QTA
👥 친구 초대 +200 QTA × 무제한
📅 매일 로그인 +10 QTA × 3회
🤝 거래 완료 +10 QTA

→ 5,000 QTA 모이면 출금!""",
    """당근에 전화번호 적기 싫어서
가지를 시작했어요.

닉네임 하나면 끝.
거래 후엔 메시지도 사라져요.
🍆 Eggplant""",
    """"가지에서 거래하면 토큰이 쌓여요"
"그게 뭔데?"
"그게 진짜 돈이래.\"""",
]
for c in short_copies:
    add_quote_box(doc, c)
    doc.add_paragraph()

add_heading(doc, '📺 영상 광고 30초 스크립트', level=2, color=DARK_PURPLE)
scenes = [
    ('Scene 1 (4초)',
     '당근에서 거래하다 받은 광고 문자 알림이 화면 가득.',
     '나레이션: "거래 한 번 하고… 광고 문자만 한 달째."'),
    ('Scene 2 (6초)',
     '스마트폰 화면이 가지(🍆) 앱으로 전환. 가입 화면에 닉네임만 입력하는 모습.',
     '나레이션: "닉네임 하나면 끝."'),
    ('Scene 3 (6초)',
     'QR 코드를 보여주고 만남. 거래 후 채팅 화면이 후룩 사라짐.',
     '나레이션: "전화번호도, 메시지 기록도 남지 않아요."'),
    ('Scene 4 (8초)',
     'QTA 잔고가 +500 → +700 → +1,000으로 올라감.',
     '나레이션: "가입만 해도 500. 친구 초대마다 200. 5,000 모이면 출금까지."'),
    ('Scene 5 (6초)',
     '🍆 로고 + 슬로건 등장.',
     '나레이션: "익명으로 안전한 중고거래, Eggplant." / 자막: 지금 다운로드 → eggplant.life'),
]
for label, action, narr in scenes:
    add_para(doc, label, size=11, bold=True, color=PURPLE)
    add_para(doc, f'  화면: {action}', size=10, color=TEXT_DARK)
    add_para(doc, f'  음성: {narr}', size=10, color=GRAY, italic=True)
    doc.add_paragraph()


# ═════ 10. SNS 가이드 ═════
add_pagebreak(doc)
add_heading(doc, '10. SNS / 광고 콘텐츠 가이드', level=1, color=PURPLE)

add_heading(doc, '📱 인스타그램 / 틱톡 (15초)', level=2, color=DARK_PURPLE)
add_bullet(doc, '후크: "당근에 전화번호 적기 싫어요?" (1초)')
add_bullet(doc, '솔루션: "닉네임만으로 시작하는 가지 🍆" (3초)')
add_bullet(doc, '차별점: "가입만 해도 +500 QTA" (3초)')
add_bullet(doc, '만남 장면: QR 보여주고 거래 (5초)')
add_bullet(doc, 'CTA: "지금 가지 다운로드" (3초)')

add_heading(doc, '🐦 트위터 / X', level=2, color=DARK_PURPLE)
add_quote_box(doc,
    '중고거래 하면서 모은 광고 문자만 247통.\n'
    '가지(@eggplant)는 닉네임만 받아요.\n'
    '가입 보너스 +500 QTA 즉시 지급.\n'
    '🍆 https://eggplant.life'
)

add_heading(doc, '📺 YouTube Shorts', level=2, color=DARK_PURPLE)
add_bullet(doc, '15초: "전화번호 없이 중고거래 가능?" → 답변')
add_bullet(doc, '30초: 가입부터 첫 거래까지 풀 시연')

doc.add_paragraph()
add_heading(doc, '🎨 비주얼 가이드', level=2, color=DARK_PURPLE)
add_bullet(doc, '메인 컬러: Eggplant Purple #9333EA')
add_bullet(doc, '포인트 컬러: Green Leaf #22C55E')
add_bullet(doc, '마스코트: 🍆 가지 캐릭터 (이미 제작됨)')
add_bullet(doc, '톤앤매너: 친근하고, 약간 장난스럽게, 그러나 신뢰감 있게')


# ═════ 11. FAQ ═════
add_pagebreak(doc)
add_heading(doc, '11. FAQ — 자주 묻는 질문', level=1, color=PURPLE)

faqs = [
    ('Q1. 정말로 전화번호 없이 가입돼요?',
     '네. 닉네임 + 지갑주소(0x로 시작하는 40자 hex) + 비밀번호만 입력하면 됩니다. '
     '전화번호 입력란 자체가 앱에 존재하지 않습니다.'),
    ('Q2. QTA가 진짜 돈이 되는 건가요?',
     '5,000 QTA 이상 모이면 출금 신청 가능합니다. 운영자 승인 후 처리됩니다. '
     '환율 정책은 별도 공지됩니다.'),
    ('Q3. 채팅 기록이 정말 안 남아요?',
     '메시지는 Cloudflare Durable Object의 메모리에만 일시적으로 존재합니다. '
     'D1(영구 DB)에 저장하지 않으며, 채팅방을 나가면 즉시 휘발됩니다.'),
    ('Q4. 사기당하면 어떡해요?',
     '① 매너온도가 낮은 사용자는 피하세요. ② QR 만남 후 직거래를 권장합니다. '
     '③ QTA 결제는 거래 완료 시 자동 송금되므로 송금 사기는 불가능합니다. '
     '④ 신고 기능으로 즉시 차단할 수 있습니다.'),
    ('Q5. 다른 앱과 뭐가 달라요?',
     '핵심: 익명 + 휘발성 + 토큰 3박자입니다. 이 모두를 결합한 마켓은 Eggplant가 유일합니다.'),
    ('Q6. iOS도 되나요?',
     '현재 Android APK 우선 출시 중입니다. iOS는 다음 버전에서 지원 예정입니다.'),
    ('Q7. 비밀번호를 잊었어요.',
     '가입 시 입력한 지갑주소로 비밀번호 재설정 가능합니다. '
     '전화번호 인증이 없으므로 지갑주소를 반드시 안전하게 보관하세요.'),
    ('Q8. 친구 초대 보상은 진짜 무제한인가요?',
     '네. 1명당 +200 QTA, 인원 제한 없음. 다만 본인의 다른 계정 자기 초대(셀프 리퍼럴)는 자동 차단됩니다.'),
    ('Q9. 거래 후 별로면 환불되나요?',
     '직거래 특성상 환불은 거래 당사자 간 합의로 처리됩니다. '
     '거래 시작 전 매너온도와 후기를 꼭 확인하세요.'),
    ('Q10. 데이터 삭제하고 탈퇴하면 진짜 다 지워지나요?',
     '네. 회원 탈퇴 시 본인의 데이터(상품·후기·잔액 등)는 즉시 영구 삭제됩니다. 백업도 남기지 않습니다.'),
]
for q, a in faqs:
    add_para(doc, q, size=12, bold=True, color=PURPLE)
    add_para(doc, a, size=11, color=TEXT_DARK)
    doc.add_paragraph()


# ═════ 마지막 페이지 ═════
add_pagebreak(doc)
add_heading(doc, '📞 컨택트', level=1, color=PURPLE)
add_bullet(doc, '공식 사이트: https://eggplant.life')
add_bullet(doc, '다운로드: https://github.com/Stevewon/carrot-market/releases/latest')
add_bullet(doc, '문의: (운영자 연락처)')

doc.add_paragraph()
add_heading(doc, '📌 부록 — 마케팅 활용 체크리스트', level=1, color=PURPLE)
checklist = [
    '공식 인스타그램 계정 개설 → 마스코트 활용 콘텐츠 주 3회',
    '틱톡 챌린지 — "#가지챌린지" + QTA 보너스 이벤트',
    '대학가 오프라인 QR 포스터 (가입 시 +1,000 QTA 추가 코드)',
    '카페·중고거래 커뮤니티 게시판 운영 가이드',
    '인플루언서 시드 — 학과 단톡방 / 맘카페 위주',
    '보도자료 발송 — 테크크런치 / 디지털타임스 / 더밀크',
    '앱스토어/플레이스토어 ASO 키워드 — "익명 중고거래", "QR 거래", "토큰 마켓"',
]
for c in checklist:
    add_bullet(doc, '☐ ' + c)

doc.add_paragraph()
doc.add_paragraph()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = final.add_run('🍆 Eggplant — 익명으로 안전한 중고거래, 사용자가 보상받는 마켓.')
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(13)
run.bold = True
run.font.color.rgb = PURPLE

final2 = doc.add_paragraph()
final2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = final2.add_run('Built with Flutter + Cloudflare Workers, Designed for Privacy.')
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(10)
run.italic = True
run.font.color.rgb = GRAY


# ═════ 저장 ═════
out = os.path.join(os.path.dirname(__file__), 'Eggplant_홍보보고서.docx')
doc.save(out)
print(f'생성 완료: {out}')
print(f'파일 크기: {os.path.getsize(out)/1024:.1f} KB')
